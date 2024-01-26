import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import g4f

def Aige(file_name):
    # Define the input for the model
    input = {"role": "user", "content": f"写一个{file_name.replace('自评','')}自评，注意生成的回答中应该具有创新性，并且以第二人称创作并且不必表明出处，直接输出结果即可，并按照下面段落的格式：韩梅梅，你乐于助人，与同学相处融洽，有着良好的人际关系。你热情大方，积极参加班级、学校的各项活动。作为班里的团支书，你工作认真负责，用心协助老师做好工作。外表阳光的你，内心极为细腻，作文课上总是能听到你优美生动的范文，成绩单上总有你骄人的成绩，愿你在广阔空间里展翅翱翔。"}

    # Use the g4f package with the gpt-4 model
    response = g4f.ChatCompletion.create(
        model=g4f.models.gpt_4_32k_0613,
        messages=[input]
    )

    # Return the output of the model
    return response

def write_to_word(file_path, file_name, content):
    # Create a new Word document
    doc = Document()

    # Add the title to the document
    title = doc.add_paragraph(file_name)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = title.runs[0]
    run.bold = True
    font = run.font
    font.size = Pt(14)

    # Add the content to the document
    doc.add_paragraph(content)

    # Save the document
    full_path = os.path.join(file_path, file_name + '.docx')
    doc.save(full_path)


# Specify the file path and name
file_path = '/Users/mac/Desktop/coding'
file_name = ['同学姓名1自评','同学姓名2自评','同学姓名3自评']
def final(file_path,file_name):
    # Call the Aige function
    output = Aige(file_name)
    # Write the output to a Word document
    write_to_word(file_path, file_name, output)

for i in file_name:
    final(file_path,i)

